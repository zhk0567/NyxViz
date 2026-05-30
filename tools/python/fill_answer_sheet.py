"""Fill official 1-II_answerSheet.docx — 上表下图, 规范中文排版."""
from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from answer_content import TASK_BUILDERS
from docx_insert import Figure, render_blocks
from export_paper_docx import load_timeline

ROOT = Path(__file__).resolve().parents[2]
STATS = ROOT / "public" / "stats" / "timeline.json"
ANSWER_MARKER = "（下面是答题区域）"
QUESTION_START = re.compile(r"^\d、")
TEMPLATE_CANDIDATES = [
    ROOT / "docs" / "competition" / "1-II_answerSheet.docx",
    ROOT / "1-II_answerSheet.docx",
]
TEAM_CONFIG = ROOT / "docs" / "competition" / "team.json"
TEAM_CONFIG_EXAMPLE = ROOT / "docs" / "competition" / "team.json.example"
DEFAULT_OUT = ROOT / "docs" / "submission" / "Nyx_answerSheet_filled.docx"
CHAR_LIMIT = 800


def find_template() -> Path:
    for p in TEMPLATE_CANDIDATES:
        if p.exists():
            return p
    raise FileNotFoundError(
        "缺少 1-II_answerSheet.docx，请放到 docs/competition/\n"
        "  https://chinavis.org/2026/challenge/1-II_answerSheet.docx"
    )


def marker_indices(doc: Document) -> list[int]:
    return [i for i, p in enumerate(doc.paragraphs) if ANSWER_MARKER in p.text]


def next_question_index(doc: Document, after: int) -> int:
    for j in range(after + 1, len(doc.paragraphs)):
        if QUESTION_START.match(doc.paragraphs[j].text.strip()):
            return j
    return len(doc.paragraphs)


def remove_paragraphs_between(doc: Document, start_exclusive: int, end_exclusive: int) -> None:
    if end_exclusive <= start_exclusive + 1:
        return
    for j in range(end_exclusive - 1, start_exclusive, -1):
        el = doc.paragraphs[j]._element
        el.getparent().remove(el)


def set_para_text(para, prefix: str, value: str) -> None:
    para.text = f"{prefix}{value}"
    for run in para.runs:
        run.font.name = "宋体"
        run.font.size = Pt(10.5)
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")


def apply_team_config(doc: Document, config: dict) -> None:
    for para in doc.paragraphs:
        t = para.text.strip()
        if t.startswith("参赛队名称：") and config.get("team_name"):
            set_para_text(para, "参赛队名称：", config["team_name"])
        elif t.startswith("团队成员：") and config.get("member_line_1"):
            set_para_text(para, "团队成员：", config["member_line_1"])
        elif t.startswith("团队成员是否与报名表一致") and config.get("consistent_with_registration"):
            para.text = f"团队成员是否与报名表一致（是或否）：{config['consistent_with_registration']}"
        elif t.startswith("是否学生队") and config.get("student_team"):
            para.text = f"是否学生队（是或否）：\t{config['student_team']}"
        elif t.startswith("使用的分析工具") and config.get("tools"):
            para.text = f"使用的分析工具或开发工具（如果使用了自己研发的软件或工具请具体说明）：{config['tools']}"
        elif t.startswith("共计耗费时间") and config.get("person_days"):
            para.text = f"共计耗费时间（人天）： {config['person_days']}"
        elif t.startswith("本次比赛结束后") and config.get("publish_ok"):
            para.text = f"本次比赛结束后，我们是否可以在网络上公布该答卷与视频（是或否）：{config['publish_ok']}"

    extras = config.get("member_lines_extra")
    if extras:
        lines = extras if isinstance(extras, list) else [extras]
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip().startswith("团队成员：") and i + 1 < len(doc.paragraphs):
                for j, line in enumerate(lines):
                    if i + 1 + j < len(doc.paragraphs):
                        nxt = doc.paragraphs[i + 1 + j].text.strip()
                        if not nxt.startswith("团队成员是否与"):
                            doc.paragraphs[i + 1 + j].text = line
                break


def fill_answer_region(
    doc: Document,
    task_index: int,
    timeline: dict,
    *,
    char_limit: int,
) -> tuple[int, int]:
    markers = marker_indices(doc)
    if task_index >= len(markers):
        raise IndexError(f"任务{task_index + 1}：未找到答题区域标记")
    marker_idx = markers[task_index]
    end_idx = next_question_index(doc, marker_idx)
    remove_paragraphs_between(doc, marker_idx, end_idx)

    marker_para: Paragraph = doc.paragraphs[marker_idx]
    blocks = TASK_BUILDERS[task_index](timeline)
    _, chars = render_blocks(marker_para, doc, blocks)
    n_fig = sum(1 for b in blocks if isinstance(b, Figure))
    if chars > char_limit:
        print(f"  警告：任务{task_index + 1} 正文约{chars}字，超过建议{char_limit}字", file=sys.stderr)
    return chars, n_fig


def fill_answer_sheet(
    template: Path,
    output: Path,
    timeline: dict,
    *,
    team_config: Path | None = None,
    char_limit: int = CHAR_LIMIT,
) -> list[tuple[int, int, int]]:
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, output)
    doc = Document(str(output))

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    cfg_path = team_config or TEAM_CONFIG
    if not cfg_path.exists() and TEAM_CONFIG_EXAMPLE.exists():
        cfg_path = TEAM_CONFIG_EXAMPLE
    if cfg_path.exists():
        apply_team_config(doc, json.loads(cfg_path.read_text(encoding="utf-8")))
        print(f"已应用队名配置：{cfg_path.name}")

    stats: list[tuple[int, int, int]] = []
    for i, _ in enumerate(TASK_BUILDERS):
        chars, n_img = fill_answer_region(doc, i, timeline, char_limit=char_limit)
        stats.append((i + 1, chars, n_img))

    doc.save(output)
    return stats


def main() -> int:
    if sys.platform == "win32" and hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")

    parser = argparse.ArgumentParser(description="Fill ChinaVis answer sheet (official template).")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUT)
    parser.add_argument("--team-config", type=Path, default=None)
    parser.add_argument("--limit", type=int, default=CHAR_LIMIT)
    args = parser.parse_args()

    if not STATS.exists():
        print("缺少 timeline.json，请先运行 npm run precompute", file=sys.stderr)
        return 1

    try:
        template = find_template()
        stats = fill_answer_sheet(
            template,
            args.output,
            load_timeline(),
            team_config=args.team_config,
            char_limit=args.limit,
        )
    except Exception as exc:
        print(f"填充失败：{exc}", file=sys.stderr)
        return 1

    print(f"已写入 {args.output}\n")
    print(f"{'题号':>4} {'字数':>6} {'图':>4}")
    for task, chars, n_img in stats:
        print(f"{task:>4} {chars:>6} {n_img:>4}")
    if not TEAM_CONFIG.exists():
        print(f"\n请编辑 docs/competition/team.json 填写真实队名/成员后重跑。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
