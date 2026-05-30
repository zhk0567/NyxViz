"""Insert paragraphs, tables and figures after an anchor in an existing docx."""
from __future__ import annotations

from dataclasses import dataclass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

from docx_format import count_text_chars, format_body_paragraph, normalize_cn, set_run_font, style_table_cell
from report_docx_shared import resolve_image


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def insert_paragraph_after_element(element, parent) -> Paragraph:
    new_p = OxmlElement("w:p")
    element.addnext(new_p)
    return Paragraph(new_p, parent)


def insert_table_after(paragraph: Paragraph, doc: Document, rows: int, cols: int):
    table = doc.add_table(rows=rows, cols=cols)
    tbl = table._tbl
    doc.element.body.remove(tbl)
    paragraph._element.addnext(tbl)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    return table


def insert_body_after(anchor: Paragraph, text: str, *, bold: bool = False, indent: bool = True) -> Paragraph:
    text = normalize_cn(text)
    para = insert_paragraph_after(anchor)
    format_body_paragraph(para, indent=indent and not bold)
    if bold:
        para.paragraph_format.first_line_indent = Pt(0)
    set_run_font(para.add_run(text), "宋体", 10.5, bold=bold)
    return para


def insert_caption_after(anchor: Paragraph, label: str, *, above: bool) -> Paragraph:
    para = insert_paragraph_after(anchor)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6 if above else 3)
    para.paragraph_format.space_after = Pt(3 if above else 6)
    set_run_font(para.add_run(normalize_cn(label)), "宋体", 10.5, bold=above)
    return para


def insert_table_block_after(
    anchor: Paragraph,
    doc: Document,
    caption: str,
    headers: list[str],
    rows: list[list[str]],
) -> Paragraph:
    anchor = insert_caption_after(anchor, caption, above=True)
    table = insert_table_after(anchor, doc, 1 + len(rows), len(headers))
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        style_table_cell(cell, bold=True)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            style_table_cell(cell)
    return insert_paragraph_after_element(table._tbl, anchor._parent)


def insert_figure_after(
    anchor: Paragraph,
    image_name: str,
    caption: str,
    width_cm: float = 13.0,
) -> Paragraph:
    path = resolve_image(image_name)
    pic_para = insert_paragraph_after(anchor)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if path:
        pic_para.add_run().add_picture(str(path), width=Cm(width_cm))
    else:
        set_run_font(pic_para.add_run(f"（缺图 {image_name}）"), "宋体", 10.5)
    return insert_caption_after(pic_para, caption, above=False)


@dataclass
class Text:
    content: str = ""
    bold: bool = False


@dataclass
class Table:
    caption: str = ""
    headers: list[str] | None = None
    rows: list[list[str]] | None = None

    def __post_init__(self):
        self.headers = self.headers or []
        self.rows = self.rows or []


@dataclass
class Figure:
    image: str = ""
    caption: str = ""
    width_cm: float = 13.0


Block = Text | Table | Figure


def render_blocks(anchor: Paragraph, doc: Document, blocks: list[Block]) -> tuple[Paragraph, int]:
    chars = 0
    for block in blocks:
        if isinstance(block, Text):
            anchor = insert_body_after(anchor, block.content, bold=block.bold)
            chars += count_text_chars(block.content)
        elif isinstance(block, Table):
            anchor = insert_table_block_after(
                anchor, doc, block.caption, block.headers, block.rows
            )
        elif isinstance(block, Figure):
            anchor = insert_figure_after(anchor, block.image, block.caption, block.width_cm)
    return anchor, chars
