"""Export ChinaVis-style work description docx (作品说明文档)."""
from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from PIL import Image

from docx_format import normalize_cn, parse_bold_runs, set_run_font
from report_docx_shared import ROOT, resolve_image
from viz_style import DOC_EMBED_DPI
from work_doc_content import (
    Block,
    CoverInfo,
    Figure,
    Heading,
    Table,
    Text,
    build_blocks,
    load_team_config,
)

STATS = ROOT / "public" / "stats" / "timeline.json"
OUT = ROOT / "docs" / "submission" / "NyxViz_作品说明文档.docx"


def load_timeline() -> dict:
    return json.loads(STATS.read_text(encoding="utf-8"))


def set_paragraph_format(paragraph, *, indent: bool = True, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    fmt.alignment = align
    if indent:
        fmt.first_line_indent = Pt(24)


def add_rich_body(doc: Document, text: str, *, indent: bool = True) -> None:
    runs = parse_bold_runs(text)
    if not runs:
        return
    p = doc.add_paragraph()
    set_paragraph_format(p, indent=indent)
    for segment, bold in runs:
        set_run_font(p.add_run(segment), "宋体", 12, bold=bold)


def add_body(doc: Document, text: str, *, indent: bool = True) -> None:
    add_rich_body(doc, text, indent=indent)


def add_heading(doc: Document, text: str, level: int) -> None:
    text = normalize_cn(text)
    sizes = {0: 18, 1: 16, 2: 14, 3: 12}
    fonts = {0: "黑体", 1: "黑体", 2: "黑体", 3: "宋体"}
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(12 if level <= 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(text), fonts.get(level, "宋体"), sizes.get(level, 12), bold=True)


def add_cover_line(doc: Document, text: str, *, center: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(3)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(normalize_cn(text)), "宋体", 12)


def add_cover_page(doc: Document, cover: CoverInfo) -> None:
    add_heading(doc, "2026 年第十三届中国可视化与可视分析大会", 0)
    add_heading(doc, "数据可视化竞赛赛道 II（ChinaVis Data Challenge 2026 - mini challenge II）", 3)
    add_heading(doc, "答  卷", 0)
    doc.add_paragraph()

    add_cover_line(doc, f"参赛队名称：{cover.team_name}")
    add_cover_line(doc, f"团队成员：  {cover.member_line_1}")
    if cover.member_lines_extra:
        for line in cover.member_lines_extra.split("\n"):
            line = line.strip()
            if line:
                add_cover_line(doc, line)
    add_cover_line(doc, f"团队成员是否与报名表一致（是或否）：{cover.consistent_with_registration}")
    add_cover_line(doc, f"是否学生队（是或否）：{cover.student_team}")
    add_cover_line(doc, f"使用的分析工具或开发工具（如果使用了自己研发的软件或工具请具体说明）：{cover.tools}")
    add_cover_line(doc, f"共计耗费时间（人天）： {cover.person_days}")
    add_cover_line(doc, f"本次比赛结束后，我们是否可以在网络上公布该答卷与视频（是或否）：{cover.publish_ok}")
    doc.add_paragraph()
    add_body(
        doc,
        "（以下为作品说明正文：系统概览、四题案例分析、综合发现与附录。配图按出现顺序编号为图1、图2…；"
        "同一配图重复引用时标注「见图N，同图从略」。子图含义见各图注。）",
        indent=False,
    )
    doc.add_page_break()


def add_table_caption(doc: Document, num: int, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    label = f"表{num} "
    set_run_font(p.add_run(label), "宋体", 10.5, bold=True)
    set_run_font(p.add_run(normalize_cn(title)), "宋体", 10.5, bold=True)


def add_figure_caption(doc: Document, num: int, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(12)
    set_run_font(p.add_run(f"图{num} "), "宋体", 10.5, bold=True)
    set_run_font(p.add_run(normalize_cn(title)), "宋体", 10.5, bold=False)


def style_table(table) -> None:
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for run in para.runs:
                    set_run_font(run, "宋体", 10.5)


def add_table_block(
    doc: Document,
    num: int,
    title: str,
    headers: list[str],
    rows: list[list[str]],
) -> None:
    add_table_caption(doc, num, title)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    style_table(table)
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_run_font(run, "宋体", 10.5, bold=True)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    set_run_font(run, "宋体", 10.5)
    doc.add_paragraph()


def figure_embed_width_cm(path: Path, max_cm: float = 16.5) -> float:
    with Image.open(path) as im:
        natural_cm = im.width / DOC_EMBED_DPI * 2.54
    return min(max_cm, natural_cm)


def add_figure_block(
    doc: Document,
    num: int,
    image_name: str,
    caption: str,
    width_cm: float,
    seen: dict[str, int],
) -> bool:
    """Embed figure; duplicate filenames become cross-refs. Returns True if newly embedded."""
    if image_name in seen:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        first = seen[image_name]
        set_run_font(p.add_run(f"（见图{first}："), "宋体", 10.5, bold=True)
        set_run_font(p.add_run(normalize_cn(caption)), "宋体", 10.5, bold=False)
        set_run_font(p.add_run("，同图从略）"), "宋体", 10.5, bold=True)
        return False

    path = resolve_image(image_name)
    if not path:
        add_body(doc, f"（缺图：{image_name}）", indent=False)
        return False
    embed_cm = figure_embed_width_cm(path, max_cm=min(16.5, width_cm))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(embed_cm))
    add_figure_caption(doc, num, caption)
    seen[image_name] = num
    return True


def setup_page(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def render_blocks(doc: Document, blocks: list[Block]) -> None:
    fig_no = 1
    tbl_no = 1
    seen_figures: dict[str, int] = {}
    for block in blocks:
        if isinstance(block, Heading):
            add_heading(doc, block.text, block.level)
        elif isinstance(block, Text):
            add_rich_body(doc, block.content, indent=block.indent)
        elif isinstance(block, Table):
            add_table_block(doc, tbl_no, block.caption, block.headers, block.rows)
            tbl_no += 1
        elif isinstance(block, Figure):
            if add_figure_block(doc, fig_no, block.file, block.caption, block.width_cm, seen_figures):
                fig_no += 1


def build_document(timeline: dict, cover: CoverInfo | None = None) -> Document:
    doc = Document()
    setup_page(doc)
    add_cover_page(doc, cover or load_team_config())
    render_blocks(doc, build_blocks(timeline))
    return doc


def main() -> int:
    if not STATS.exists():
        print("Missing timeline.json — run: npm run precompute", file=sys.stderr)
        return 1

    timeline = load_timeline()
    doc = build_document(timeline)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
