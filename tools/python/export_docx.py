"""Export submission docx from report markdown and figures."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[2]
REPORT = ROOT / "docs" / "report"
FIGURES = ROOT / "docs" / "figures"
OUT = ROOT / "docs" / "report" / "Nyx_Submission.docx"

SECTIONS = [
    ("task1_volume.md", "任务一：体数据渲染与密度演化"),
    ("task2_evolution.md", "任务二：宇宙密度演化规律归纳"),
    ("task3_histogram.md", "任务三：时序密度对数直方图统计"),
    ("task4_brush.md", "任务四：相空间交互刷选可视分析"),
]

TASK1_IMAGES = [
    "task1_vol_strip.png",
    "task1_vol_t0000.png",
    "task1_vol_t0050.png",
    "task1_vol_t0099.png",
]

TASK3_IMAGES = [
    "task3_hist_overlay.png",
    "task3_metrics_timeline.png",
    "task3_evolution_metrics.png",
]
TASK4_IMAGES = [
    "task4_brush_triptych.png",
    "task4_hist_brush_top1.png",
    "task4_brush_top1.png",
]


def resolve_image(name: str) -> Path | None:
    candidates = [
        name,
        name.replace("task1_vol_", "task1_"),
        name.replace("task1_vol_", "task1_slice_"),
    ]
    for c in candidates:
        p = FIGURES / c
        if p.exists():
            return p
    return None


def md_to_paragraphs(text: str) -> list[tuple[str, str]]:
    """Return list of (type, content): heading2, heading3, bullet, para."""
    blocks: list[tuple[str, str]] = []
    for line in text.splitlines():
        if line.startswith("## "):
            blocks.append(("heading2", line[3:].strip()))
        elif line.startswith("### "):
            blocks.append(("heading3", line[4:].strip()))
        elif line.startswith("- "):
            blocks.append(("bullet", line[2:].strip()))
        elif line.startswith("![") or not line.strip():
            continue
        else:
            blocks.append(("para", line.strip()))
    return blocks


def add_section(doc: Document, md_path: Path, extra_images: list[str]) -> None:
    if not md_path.exists():
        return
    text = md_path.read_text(encoding="utf-8")
    for kind, content in md_to_paragraphs(text):
        if kind == "heading2":
            doc.add_heading(content, level=1)
        elif kind == "heading3":
            doc.add_heading(content, level=2)
        elif kind == "bullet":
            doc.add_paragraph(content, style="List Bullet")
        else:
            doc.add_paragraph(content)

    for img_name in extra_images:
        img_path = resolve_image(img_name)
        if img_path:
            doc.add_paragraph(img_path.stem)
            doc.add_picture(str(img_path), width=Cm(15))


def main() -> int:
    if not REPORT.exists():
        print(f"Missing {REPORT}", file=sys.stderr)
        return 1

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)

    doc.add_heading("Nyx 赛题 II 科学可视化提交材料", 0)
    doc.add_paragraph(
        "基于 Nyx 128³ 气体密度场（100 时间步）的体渲染、时序统计与刷选联动分析。"
    )

    extras = {
        "task1_volume.md": TASK1_IMAGES,
        "task2_evolution.md": [
            "task2_evolution_story.png",
            "task3_hist_overlay.png",
            "task1_vol_t0000.png",
            "task1_vol_t0099.png",
        ],
        "task3_histogram.md": TASK3_IMAGES,
        "task4_brush.md": TASK4_IMAGES,
    }

    for md_name, _title in SECTIONS:
        add_section(doc, REPORT / md_name, extras.get(md_name, []))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
